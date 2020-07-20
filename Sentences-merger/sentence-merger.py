import docx
import os

os.getcwd()
doc1 = docx.Document('french.docx')
doc2 = docx.Document('english.docx')
doc3 = docx.Document()
sentences = []
sentences2 = []

for i in doc1.paragraphs:
    sentences.append(i.text.split('. '))

for i in doc2.paragraphs:
    sentences2.append(i.text.split('. '))

n=0
for i in range(len(sentences)):

	for a in range(len(sentences[n])):
		print(sentences[i][a])
		french = doc3.add_paragraph().add_run(sentences[i][a])
		french.bold = True
		print(sentences2[i][a])
		print(" ")
		spanish = doc3.add_paragraph().add_run(sentences2[i][a])
		spanish.italic = True
	print(" ")
	spanish = doc3.add_paragraph().add_run(" ")
	n+=1

doc3.save('new.docx')