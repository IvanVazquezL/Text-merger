import googletrans
from googletrans import Translator
import docx
import os

os.getcwd()
translator = Translator()
doc1 = docx.Document('frenchOg.docx')
doc2 = docx.Document()
doc3 = docx.Document()
doc4 = docx.Document()
para = []
paraSpa = []
sentences = []
sentences2 = []
removeWord="Ngofeen"

for i in range(len(doc1.paragraphs)):
	if removeWord not in doc1.paragraphs[i].text:
		para.append([doc1.paragraphs[i].text])

for i in range(len(para)):
	print(i,") ",para[i])
	print(" ")

answer = input("Do you want to use the manual Remover?")

if (answer=="y"):
	nums=input("Type paragraph number")
	remover = [i for i in nums.split(",")]

	#To convert the list of strings into list of integers
	remover = list(map(int, remover))
	remover2 = remover.copy()

	for i in range(len(remover)):
		if len(remover2)>= i:
			number= remover2[i]-i
			para.pop(number)
		else:
			break

para = [[x.replace("Elodie:\xa0","") for x in l] for l in para]

for i in range(len(para)):
	print(i,") ",para[i])
	print(" ")

for i in range(len(para)):
	doc2.add_paragraph().add_run(para[i])

for i in range(len(para)):
	paraSpa.append([(translator.translate(*para[i], dest="es").text)])
	doc3.add_paragraph().add_run(*paraSpa[i])

for i in range(len(para)):
	french = doc4.add_paragraph().add_run(para[i])
	french.bold = True
	spanish = doc4.add_paragraph().add_run(*paraSpa[i])
	spanish.italice = True
	spanish = doc4.add_paragraph().add_run(" ")


doc2.save('french.docx')
doc3.save('spanishTrans.docx')
doc4.save('newTrans.docx')